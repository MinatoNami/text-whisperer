"""Rendering a Markdown summary as Word."""

from pathlib import Path

import pytest

from telegram_stt.archive import Archive
from telegram_stt.docx_export import summary_to_docx

SUMMARY = """## Summary
We agreed to ship the trailer by **Aug 31**.

## Key points
- **Batch 1:** goes to Hogarth by Aug 24
- Oguri banners are capped at 1KB
* An asterisk bullet counts too

## Decisions
None recorded.
"""


@pytest.fixture
def built(tmp_path):
    from docx import Document

    path = summary_to_docx(SUMMARY, tmp_path / "out.docx", title="71 Robinson Rd 21",
                           recorded="19 Aug 2026, 11:34", duration="51 min", language="English")
    return path, Document(str(path))


def styles(document):
    return [p.style.name for p in document.paragraphs if p.text.strip()]


def text_of(document):
    return "\n".join(p.text for p in document.paragraphs)


class TestStructure:
    def test_a_file_is_produced(self, built):
        path, _ = built
        assert path.suffix == ".docx" and path.stat().st_size > 5000

    def test_the_title_uses_a_real_title_style(self, built):
        _, document = built
        assert document.paragraphs[0].style.name == "Title"
        assert document.paragraphs[0].text == "71 Robinson Rd 21"

    def test_sections_become_real_headings(self, built):
        _, document = built
        headings = [p.text for p in document.paragraphs if p.style.name.startswith("Heading")]
        assert headings == ["Summary", "Key points", "Decisions"]

    def test_bullets_use_the_list_style_not_literal_characters(self, built):
        """A literal bullet is text; Word cannot indent or renumber it."""
        _, document = built
        assert styles(document).count("List Bullet") == 3
        assert "•" not in text_of(document)
        assert "- " not in text_of(document)

    def test_asterisk_bullets_are_recognised_too(self, built):
        _, document = built
        assert any("asterisk bullet" in p.text for p in document.paragraphs
                   if p.style.name == "List Bullet")

    def test_bold_is_character_formatting_not_asterisks(self, built):
        _, document = built
        assert "**" not in text_of(document)
        bolded = {r.text for p in document.paragraphs for r in p.runs if r.bold}
        assert "Aug 31" in bolded
        assert "Batch 1:" in bolded

    def test_the_subtitle_carries_the_context(self, built):
        _, document = built
        assert "51 min" in document.paragraphs[1].text
        assert "English" in document.paragraphs[1].text

    def test_no_paragraph_contains_a_newline(self, built):
        """Word has no in-paragraph newlines; each line must be its own."""
        _, document = built
        assert all("\n" not in p.text for p in document.paragraphs)

    def test_it_is_a_valid_ooxml_package_with_numbering(self, built):
        import zipfile

        path, _ = built
        with zipfile.ZipFile(path) as archive:
            names = archive.namelist()
        for part in ("word/document.xml", "word/styles.xml", "word/numbering.xml"):
            assert part in names, f"{part} missing — Word will not open this"


class TestEdgeCases:
    def test_empty_markdown_still_produces_a_readable_document(self, tmp_path):
        from docx import Document

        path = summary_to_docx("", tmp_path / "empty.docx", title="Nothing")
        assert Document(str(path)).paragraphs[0].text == "Nothing"

    def test_missing_context_omits_the_subtitle(self, tmp_path):
        from docx import Document

        path = summary_to_docx("Body text.", tmp_path / "bare.docx", title="T")
        document = Document(str(path))
        assert "·" not in document.paragraphs[1].text

    def test_unmatched_bold_markers_are_left_alone(self, tmp_path):
        from docx import Document

        path = summary_to_docx("A **dangling marker", tmp_path / "odd.docx", title="T")
        assert "**dangling" in "\n".join(p.text for p in Document(str(path)).paragraphs)


class TestFromTheArchive:
    def test_summary_docx_renders_from_the_stored_markdown(self, tmp_path, tone_wav):
        from docx import Document

        archive = Archive(tmp_path / "archive")
        entry = archive.store(
            chat_id=-1, message_id=1, source_audio=tone_wav, transcript_text="[00:00] hi",
            media_kind="audio", original_name="Board meeting.m4a", language="en",
            model="m", audio_seconds=3077.0, elapsed_seconds=29.0,
            segments=[{"start": 0.0, "end": 1.0, "text": "hi"}])
        record = archive.records()[0]

        assert archive.summary_docx(record, tmp_path / "a.docx") is None, "nothing to render yet"

        archive.write_summary(record, SUMMARY)
        path = archive.summary_docx(record, tmp_path / "b.docx")
        document = Document(str(path))
        assert document.paragraphs[0].text == "Board meeting"
        assert "51 min" in document.paragraphs[1].text
        assert "English" in document.paragraphs[1].text
